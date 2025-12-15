from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from typing import List
from io import BytesIO
import os
# from groq from Groq
from PyPDF2 import PdfReader
from docx import Document
from dotenv import load_dotenv

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.memory import ConversationBufferMemory
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.chains import ConversationalRetrievalChain
from langchain_groq import ChatGroq

load_dotenv()

app = FastAPI()

# ---- CORS (React access) ----
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---- Global variables ----
vectorstore = None
conversation = None


# ---------- TEXT EXTRACTION ----------

def extract_pdf(content: bytes):
    reader = PdfReader(BytesIO(content))
    text = ""
    for page in reader.pages:
        t = page.extract_text()
        if t:
            text += t
    return text


def extract_docx(content: bytes):
    doc = Document(BytesIO(content))
    return "\n".join(p.text for p in doc.paragraphs)


def extract_txt(content: bytes):
    return content.decode("utf-8", errors="ignore")


# ---------- UPLOAD API ----------

@app.post("/upload")
async def upload_documents(files: List[UploadFile] = File(...)):
    global vectorstore, conversation

    full_text = ""

    for file in files:
        content = await file.read()

        if file.filename.endswith(".pdf"):
            full_text += extract_pdf(content)

        elif file.filename.endswith(".docx"):
            full_text += extract_docx(content)

        elif file.filename.endswith(".txt"):
            full_text += extract_txt(content)

    if not full_text.strip():
        raise HTTPException(status_code=400, detail="No readable text found")

    # Split text
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = splitter.split_text(full_text)

    # Embeddings + FAISS
    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2"
    )
    vectorstore = FAISS.from_texts(chunks, embeddings)

    # LLM
    llm = ChatGroq(
        groq_api_key=os.getenv("GROQ_API_KEY"),
        model_name="llama-3.1-8b-instant"
    )

    conversation = ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=vectorstore.as_retriever(),
        memory=ConversationBufferMemory(
                 memory_key="chat_history",
                 return_messages=True
               )
    )

    return {"status": "success", "chunks": len(chunks)}


# ---------- CHAT API ----------

@app.post("/ask")
async def ask_question(req: dict):
    global conversation

    if conversation is None:
        raise HTTPException(status_code=400, detail="Upload documents first")

    question = req.get("question")

    response = conversation.invoke({
        "question": question
    })

    return {"answer": response["answer"]}















# import streamlit as st
# import os
# from dotenv import load_dotenv
# from io import BytesIO
# from PyPDF2 import PdfReader
# from groq import Groq

# from langchain_community.vectorstores import FAISS
# from langchain_community.embeddings import HuggingFaceEmbeddings
# from docx import Document
# from langchain.text_splitter import CharacterTextSplitter
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain.memory import ConversationBufferMemory
# from langchain.chains import ConversationalRetrievalChain

# from langchain_groq import ChatGroq


# load_dotenv()


# def get_pdf_text(pdf_docs):
#     text = ""
#     for pdf in pdf_docs:
#         pdf_bytes = pdf.read()
#         reader = PdfReader(BytesIO(pdf_bytes))
#         for page in reader.pages:
#             content = page.extract_text()
#             if content:
#                 text += content
#         pdf.seek(0)
#     return text

# def get_txt_text(txt_docs):
#     text = ""
#     for txt in txt_docs:
#         if hasattr(txt, "read"):
#             raw = txt.read()
#             txt.seek(0)

#         else:
#             raw = txt

#         try:
#             text += raw.decode("utf-8") + "\n"
#         except:
#             text += raw.decode("latin-1") + "\n"  
#     return text
# def get_docx_text(docx_docs):
#     text = ""
#     for docx in docx_docs:
#         document = Document(docx)
#         for para in document.paragraphs:
#             if para.text:
#                 text += para.text + "\n"
#         docx.seek(0)
#     return text


# def extract_all_text(pdf_docs, docx_docs, txt_docs):
#     text = ""
#     if pdf_docs:
#         text += get_pdf_text(pdf_docs)
#     if docx_docs:
#         for d in docx_docs:
#             text += get_docx_text(d)
#     if txt_docs:
#         for t in txt_docs:
#             text += get_txt_text(t)
#     return text


# def get_text_chunks(text):
#     splitter = RecursiveCharacterTextSplitter(
#         chunk_size=1000,
#         chunk_overlap=200,
#         separators=["\n"]
#     )
#     return splitter.split_text(text)

# def get_vectorstore(text_chunks):
#     embeddings = HuggingFaceEmbeddings(
#         model_name="sentence-transformers/all-MiniLM-L6-v2"
#     )
#     return FAISS.from_texts(text_chunks, embedding=embeddings)


# def get_conversation_chain(vectorstore):
#     memory = ConversationBufferMemory(
#         memory_key="chat_history",
#         return_messages=True
#     )

#     llm = ChatGroq(
#     groq_api_key=os.getenv("GROQ_API_KEY"),
#     model_name="llama-3.1-8b-instant"
#     )

#     return ConversationalRetrievalChain.from_llm(
#         llm=llm,
#         retriever=vectorstore.as_retriever(),
#         memory=memory
#     )


# def main():
#     st.set_page_config(page_title="Chat with Documents", page_icon="📚")
#     st.header("📚 Chat with Documents")

#     if "conversation" not in st.session_state:
#         st.session_state.conversation = None

#     user_question = st.text_input("Ask a question about your documents...")
#     if user_question and st.session_state.conversation:
#         response = st.session_state.conversation({"question": user_question})
#         st.write(response["answer"])

#     with st.sidebar:
#         st.subheader("Upload Your Documents")

#         pdf_docs = st.file_uploader("Upload PDF files", type="pdf", accept_multiple_files=True)
#         docx_docs = st.file_uploader("Upload Word files", type="docx", accept_multiple_files=True)
#         txt_docs = st.file_uploader("Upload Text files", type="txt", accept_multiple_files=True)
#         if st.button("Process"):
#             if not pdf_docs and not docx_docs and not txt_docs:
#                 st.error("Please upload at least one file!")
#             else:
#                 with st.spinner("Extracting text..."):
#                     text = extract_all_text(pdf_docs,docx_docs,txt_docs)

#                 with st.spinner("Splitting text..."):
#                     chunks = get_text_chunks(text)

#                 with st.spinner("Creating vector database..."):
#                     vectorstore = get_vectorstore(chunks)

#                 with st.spinner("Preparing model..."):
#                     st.session_state.conversation = get_conversation_chain(vectorstore)

#                 st.success("Document processed successfully! 🚀")

# if __name__ == "__main__":
#     main()
